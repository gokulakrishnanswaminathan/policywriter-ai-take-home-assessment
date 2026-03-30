"""Shared test fixtures for the resume parser test suite."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock

import pytest


@pytest.fixture
def sample_resume_text() -> str:
    """A realistic resume text snippet for testing extractors."""
    return (
        "Jane Doe\n"
        "jane.doe@gmail.com\n"
        "+1 (555) 123-4567\n"
        "San Francisco, CA\n\n"
        "Senior Software Engineer\n\n"
        "Summary\n"
        "Experienced engineer with 8+ years in backend development.\n\n"
        "Skills\n"
        "Python, Machine Learning, Docker, AWS, PostgreSQL, REST APIs\n\n"
        "Experience\n"
        "Tech Corp — Senior Engineer (2020–Present)\n"
        "- Designed and deployed ML pipelines using Python and TensorFlow.\n"
        "- Managed AWS infrastructure for a SaaS platform.\n\n"
        "Education\n"
        "B.S. Computer Science, Stanford University\n"
    )


@pytest.fixture
def minimal_resume_text() -> str:
    """Minimal resume with just name and email."""
    return "John Smith\njohn.smith@example.com\n"


@pytest.fixture
def no_email_resume_text() -> str:
    """Resume text with no email address."""
    return (
        "Alice Johnson\n"
        "Phone: 555-123-4567\n\n"
        "Skills: Python, Java, SQL\n"
    )


@pytest.fixture
def empty_text() -> str:
    """Empty string for edge case testing."""
    return ""


@pytest.fixture
def tmp_pdf(tmp_path: Path) -> Path:
    """Create a minimal valid PDF file for testing."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pdf_path = tmp_path / "test_resume.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.drawString(72, 720, "Jane Doe")
    c.drawString(72, 700, "jane.doe@gmail.com")
    c.drawString(72, 680, "Skills: Python, Machine Learning, Docker")
    c.save()
    return pdf_path


@pytest.fixture
def tmp_pdf_multipage(tmp_path: Path) -> Path:
    """Create a multi-page PDF for testing."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pdf_path = tmp_path / "multipage_resume.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.drawString(72, 720, "Jane Doe")
    c.drawString(72, 700, "jane.doe@gmail.com")
    c.showPage()
    c.drawString(72, 720, "Skills: Python, Docker, AWS")
    c.save()
    return pdf_path


@pytest.fixture
def tmp_empty_pdf(tmp_path: Path) -> Path:
    """Create an empty PDF (no text content) for edge case testing."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pdf_path = tmp_path / "empty_resume.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.showPage()
    c.save()
    return pdf_path


@pytest.fixture
def tmp_docx(tmp_path: Path) -> Path:
    """Create a minimal valid Word document for testing."""
    import docx

    docx_path = tmp_path / "test_resume.docx"
    doc = docx.Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane.doe@gmail.com")
    doc.add_paragraph("Skills: Python, Machine Learning, Docker")
    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def tmp_docx_with_table(tmp_path: Path) -> Path:
    """Create a Word document with a table for testing."""
    import docx

    docx_path = tmp_path / "table_resume.docx"
    doc = docx.Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane.doe@gmail.com")

    # Skills in a table (common resume layout)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "Machine Learning"
    table.cell(1, 0).text = "Docker"
    table.cell(1, 1).text = "AWS"
    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def tmp_empty_docx(tmp_path: Path) -> Path:
    """Create an empty Word document for edge case testing."""
    import docx

    docx_path = tmp_path / "empty_resume.docx"
    doc = docx.Document()
    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def mock_gemini_response() -> MagicMock:
    """Create a mock Gemini API response with a JSON skills array."""
    response = MagicMock()
    response.text = '["Python", "Machine Learning", "Docker", "AWS"]'
    return response


@pytest.fixture
def mock_gemini_response_markdown() -> MagicMock:
    """Create a mock Gemini response wrapped in markdown code fences."""
    response = MagicMock()
    response.text = '```json\n["Python", "Machine Learning", "Docker"]\n```'
    return response
