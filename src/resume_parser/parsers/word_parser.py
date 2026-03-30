"""Word document (.docx) parser using python-docx."""

from __future__ import annotations

from resume_parser.exceptions import ParsingError
from resume_parser.parsers.base import FileParser


class WordParser(FileParser):
    """Extracts text from Word (.docx) files using python-docx.

    Reads all paragraphs and table cells, preserving the document's
    reading order.
    """

    def parse(self, file_path: str) -> str:
        """Extract text from a Word document.

        Args:
            file_path: Path to the .docx file.

        Returns:
            Concatenated text from paragraphs and tables.

        Raises:
            FileNotFoundError: If the file does not exist.
            ParsingError: If the document cannot be read.
        """
        path = self._validate_file_exists(file_path)

        try:
            import docx
        except ImportError as exc:
            raise ParsingError(
                "python-docx is required to parse Word files. "
                "Install it with: pip install python-docx"
            ) from exc

        try:
            document = docx.Document(str(path))

            # Extract paragraph text
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            self.logger.info(
                "Found %d non-empty paragraph(s) in '%s'",
                len(paragraphs),
                path.name,
            )

            # Extract table cell text (resumes often use tables for layout)
            table_texts: list[str] = []
            for table in document.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        table_texts.append(" | ".join(row_cells))

            if table_texts:
                self.logger.info(
                    "Extracted text from %d table row(s) in '%s'",
                    len(table_texts),
                    path.name,
                )

            all_text = paragraphs + table_texts
            full_text = "\n".join(all_text)
            self.logger.info(
                "Extracted %d characters from '%s'", len(full_text), path.name
            )
            return full_text

        except Exception as exc:
            if isinstance(exc, (FileNotFoundError, ParsingError)):
                raise
            raise ParsingError(
                f"Failed to parse Word file '{path.name}': {exc}"
            ) from exc
